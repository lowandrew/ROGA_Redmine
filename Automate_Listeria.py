# This is where I'll learn to use python-docx so that ROGAs can get automated.

from Automate import Automate


class Listeria_Automator(Automate):

    # This prints the GeneSeekr Analysis table and the Sequence Data Quality Table to the doc.
    def print_tables_to_docx(self):
        import docx
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = docx.Document("listeria_template.docx")
        j = 2
        # GeneSeekr Analysis Table.
        for name in self.names:
            # Add a new row to the table for each SEQID to be added.
            doc.tables[2].add_row()
            # Input sample names.
            Automate.add_text_to_cell(self.metadata[name]["TextualID"] + ""
                        "\n(" + self.metadata[name]["IsolateID"] + ")", j, 0, doc.tables[2])
            # doc.tables[2].cell(j, 0).paragraphs[0].add_run(self.metadata[name]["TextualID"] + ""
            #            "\n(" + self.metadata[name]["IsolateID"] + ")").alignment = WD_ALIGN_PARAGRAPH.CENTER
            # doc.tables[2].cell(j, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            Automate.add_text_to_cell("L. monocytogenes", j, 1, doc.tables[2])
            Automate.add_text_to_cell(self.metadata[name]["Source"], j, 2, doc.tables[2])
            Automate.add_text_to_cell(self.metadata[name]["1Enzyme"], j, 3, doc.tables[2])
            Automate.add_text_to_cell(self.metadata[name]["2Enzyme"], j, 4, doc.tables[2])
            # Input the Listeria specific marker genes.
            # This way of doing things is really inflexible and generally not great, but it'll work for now.
            for i in range(6, 9):
                # If the gene name was found by geneSeekr, put in a bullet, otherwise put a ND.
                if 'hylA' in self.metadata[name]['geneSeekr']:
                    self.metadata[name]['geneSeekr'].append('hlyA')
                if doc.tables[2].cell(1, i).text in self.metadata[name]["geneSeekr"]:
                    Automate.add_text_to_cell(u"\u2022", j, i, doc.tables[2])
                else:
                    Automate.add_text_to_cell("ND", j, i, doc.tables[2])
            # Input MLST types.
            # Enter the MLST type.
            Automate.add_text_to_cell(self.metadata[name]["MLST"], j, 10, doc.tables[2])
            # Enter the type for each of the seven genes. Again, this code is inflexible and not great, but works.
            for i in range(11, 18):
                Automate.add_text_to_cell(self.metadata[name]["mlst_info"][i - 11], j, i, doc.tables[2])
            # For some reason there are extra (blank) cells in the table. This gets rid of them.
            doc.tables[2].cell(j, 6).merge(doc.tables[2].cell(j, 5))
            doc.tables[2].cell(j, 10).merge(doc.tables[2].cell(j, 9))
            Automate.add_text_to_cell("", j, 10, doc.tables[2])
            Automate.add_text_to_cell("", j, 6, doc.tables[2])
            j += 1

        # Sequence Data quality table.
        j = 1
        for name in self.names:
            # Add a new row to the table for each SEQID to be added.
            doc.tables[3].add_row()
            # Input SEQIDS.
            Automate.add_text_to_cell(self.metadata[name]["TextualID"] + ""
                            "\n" + "(" + self.metadata[name]["IsolateID"] + ")", j, 0, doc.tables[3])
            Automate.add_text_to_cell(name, j, 1, doc.tables[3])
            # Input total length.
            length = int(self.metadata[name]["TotalLength"])
            Automate.add_text_to_cell(str(length), j, 2, doc.tables[3])
            # Input coverage.
            cov = float(self.metadata[name]["Coverage"])
            covstr = "%.1f" % cov
            Automate.add_text_to_cell(covstr, j, 3, doc.tables[3])
            # Input % Identity GDCS, rST, and Pass/Fail.
            percent_id = float(self.metadata[name]["Matches"])/53.0 * 100.0
            Automate.add_text_to_cell(str(percent_id), j, 4, doc.tables[3])
            Automate.add_text_to_cell(self.metadata[name]["rST"], j, 5, doc.tables[3])
            if percent_id == 100.0:
                Automate.add_text_to_cell("Pass", j, 6, doc.tables[3])
            else:
                Automate.add_text_to_cell("Fail", j, 6, doc.tables[3])
            # Finally, input pipeline version.
            Automate.add_text_to_cell(self.metadata[name]["PipelineVersion"], j, 7, doc.tables[3])
            j += 1
        doc.save(self.outfile)

    # This takes care of printing the SNVPhyl summary to docx.
    def print_summary_to_docx(self):
        import docx
        doc = docx.Document(self.outfile)
        tables = doc.tables
        # TODO: Need to figure out how to do all the SNVPhyl-ing and whatnot (and if it even needs to be done.)
        tables[1].cell(0, 0).text = ""
        para = tables[1].cell(0, 0).add_paragraph("")
        para.add_run("Identification Summary:\n\n").bold = True
        if len(self.names) == 1:
            para.add_run("Strain " + self.metadata[self.names[0]]["TextualID"] + " was submitted for whole-genome sequencing and confirmed to be ")
        else:
            para.add_run(str(len(self.names)))
            para.add_run(" strains ")
            para.add_run(" (see Table 1) were submitted for whole-genome sequencing and confirmed to be ")
        para.add_run("Listeria monocytogenes").italic = True
        para.add_run(" based on the detection of probe sequences (e-probes) indicating the presence of IGS, ")
        para.add_run("hlyA, ").italic = True
        para.add_run(" and/or ")
        para.add_run("inlJ").italic = True
        para.add_run(" genes.\n\n")
        para.add_run("Multilocus sequence typing (MLST):\n\n").underline = True
        if len(self.mlstdict) > 1:
            para.add_run("\u2022 Isolates were sequence types ")
            i = 1
            for st in self.mlstdict:
                para.add_run("ST-")
                para.add_run(st)
                if i == len(self.mlstdict) - 1:
                    para.add_run(" and ")
                elif i == len(self.mlstdict):
                    para.add_run(".\n")
                else:
                    para.add_run(", ")
                i += 1
        else:
            para.add_run("\u2022 All isolates were sequence type ")
            for st in self.mlstdict:
                para.add_run("ST-")
                para.add_run(st + ".\n")

        for st in self.matching_mlst:
            if len(self.matching_mlst[st]) > 5:
                para.add_run("\tIsolates with ST-" + st + " are commonly recovered from the CFIA's food testing"
                                                          " program (2009-present).\n")
            elif len(self.matching_mlst[st]) == 1:
                para.add_run("\tIsolates with ST-" + st + " have not previously been recovered by the CFIA's"
                                                          " food testing program (2009-present).\n")
            else:
                para.add_run("\tIsolates with ST-" + st + " have previously been recovered from the CFIA's food "
                                                                  "testing program (2009-present).\n")

        para.add_run("\n")
        para.add_run("Quality Control Analysis:\n").bold = True
        if "122" not in self.matching_mlst and "85" not in self.matching_mlst:
            if len(self.metadata) == 1:
                para.add_run("\t This isolate does not match CFIA ")
            else:
                para.add_run("\t These isolates do not match CFIA ")
            para.add_run("Listeria monocytogenes ").italic = True
            para.add_run( "control strains (MLST: ST-122 and ST-85).")
        else:
            for strain in self.metadata:
                if self.metadata[strain]["MLST"] == "85":
                    para.add_run("\t\u2022 Isolate " + self.metadata[strain]["TextualID"] + " matches the ")
                    para.add_run("Listeria").italic = True
                    para.add_run(" control strain used (ST-85) in CFIA food testing laboratories.")
                elif self.metadata[strain]["MLST"] == "122":
                    para.add_run("\t\u2022 Isolate " + self.metadata[strain]["TextualID"] + " matches the ")
                    para.add_run("Listeria").italic = True
                    para.add_run(" control strain used (ST-122) in CFIA food testing laboratories.")
        doc.save(self.outfile)

    def print_references(self):
        import docx
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = docx.Document(self.outfile)
        para = doc.add_paragraph("")
        para.add_run("\n\nReferences:\n").bold = True
        para.add_run("\t1.  Xue JZ, Blais BW, Pagotto F (2010) Cloth-based hybridization array system for the"
                     " identification of foodborne ")
        para.add_run("Listeria").italic = True
        para.add_run(" and confirmation of ")
        para.add_run("Listeria monocytogenes.").italic = True
        para.add_run(" Int J. Food Safety. 12: 87-94.\n")
        para.add_run("\t2.  Ragon M, Wirth T, Hollandt F, Lavernir R, Lecuit M, et al. (2008) A new perspection on")
        para.add_run(" Listeria monocytogenes ").italic = True
        para.add_run("evolution. PLoS Pathog 4:e1000146\n")
        para.add_run("\t3.  Jolley KA, Bliss CM, Bennett JS, Bratcher HB, et al. (2012) Ribosomal multilocus sequence"
                     " typing: universal characterization of bacteria from domain to strain. Microbiology. 158:1005-15")
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run("--End Of Report--").bold = True
        doc.save(self.outfile)

    def __init__(self, args):
        self.seq_id_list = args.seq_id_list
        self.outfile = args.outfile_name
        self.nasmnt = args.nasmnt
        self.names = dict()
        self.metadata = dict()
        self.matching_mlst = dict()
        self.parse_metadata()
        self.find_matching_mlst()
        self.mlstdict = dict()
        self.mlst_to_dict()
        self.print_tables_to_docx()
        self.print_summary_to_docx()
        self.print_references()
        self.first_table()
        self.change_fonts()
        # self.add_redmine()


if __name__ == "__main__":

    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("seq_id_list")
    parser.add_argument("outfile_name")
    parser.add_argument("-n", "--nasmnt", type=str, default="/mnt/nas/", help="Where your NAS is mounted. Default is "
                                                                              "/mnt/nas/, trailing slash must be"
                                                                              "included.")
    arguments = parser.parse_args()
    x = Listeria_Automator(arguments)

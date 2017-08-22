# This is where I'll learn to use python-docx so that ROGAs can get automated.

import sys
import datetime

class Automate:
    """
    Main automation class, with common functions for different species ROGAs. Inherited by Automate_Salmonella,
    Automate_Listeria, and Automate_VTEC.
    """
    # This is what looks through the metadata to get all the information needed.
    # Looks through combinedMetadata, as well as rmlst and mlst specific data to get what's needed.
    def parse_metadata(self):
        """
        Looks through combinedMetadata, rMLST and MLST data to get most of the information needed to write the ROGA.
        Stores the information in the self.metadata dictionary.
        """
        import openpyxl
        import csv
        f = open(self.seq_id_list)
        self.names = f.readlines()
        f.close()
        num_samples = len(self.names)
        for i in range(len(self.names)):
            self.names[i] = self.names[i].replace("\n", "")
        # Go through the combined metadata file - it has most of the data we need.
        metadata = csv.DictReader(open(self.nasmnt + "WGSspades/reports/combinedMetadata.csv"))
        metadata_count = 0
        for row in metadata:
            # There has to be a more elegant way to do this.
            if row["SampleName"] in self.names:
                data = dict()
                data["Investigator"] = row["Investigator"]
                data["Coverage"] = row["AverageCoverageDepth"]
                data["TotalLength"] = row["TotalLength"]
                data["rST"] = row["rMLSTsequenceType"]
                data["PipelineVersion"] = row["PipelineVersion"]
                data["MLST"] = row["MLSTsequencetype"]
                data["geneSeekr"] = row["geneSeekrProfile"].split(";")
                self.metadata[row["SampleName"]] = data
                metadata_count += 1
        # Need to look in external WGS spades as well.
        metadata = csv.DictReader(open(self.nasmnt + "External_WGSspades/reports/combinedMetadata.csv"))
        for row in metadata:
            # There has to be a more elegant way to do this.
            if row["SampleName"] in self.names:
                data = dict()
                data["Investigator"] = row["Investigator"]
                data["Coverage"] = row["AverageCoverageDepth"]
                data["TotalLength"] = row["TotalLength"]
                data["rST"] = row["rMLSTsequenceType"]
                data["PipelineVersion"] = row["PipelineVersion"]
                data["MLST"] = row["MLSTsequencetype"]
                data["geneSeekr"] = row["geneSeekrProfile"].split(";")
                self.metadata[row["SampleName"]] = data
                metadata_count += 1



        # Also need to go through the rMLST file to make sure that all rMLST genes are covered.
        rMLST_data = csv.DictReader(open(self.nasmnt + "WGSspades/reports/rmlst.csv"))
        metadata_count = 0
        for row in rMLST_data:
            if row["Strain"] in self.names:
                self.metadata[row["Strain"]]["Matches"] = row["Matches"]
                metadata_count += 1
        # Check external runs.
        rMLST_data = csv.DictReader(open(self.nasmnt + "External_WGSspades/reports/rmlst.csv"))
        for row in rMLST_data:
            if row["Strain"] in self.names:
                self.metadata[row["Strain"]]["Matches"] = row["Matches"]



        # Finally, need to get info on the MLST sequence type.
        metadata_count = 0
        mlst_data = csv.DictReader(open(self.nasmnt + "WGSspades/reports/mlst.csv"))
        for row in mlst_data:
            if row["Strain"] in self.names:
                mlst = list()
                for i in range(1, 8):
                    mlst.append(row[str(i)])
                self.metadata[row["Strain"]]["mlst_info"] = mlst
                metadata_count += 1

        # Also from External.
        mlst_data = csv.DictReader(open(self.nasmnt + "External_WGSspades/reports/mlst.csv"))
        for row in mlst_data:
            if row["Strain"] in self.names:
                mlst = list()
                for i in range(1, 8):
                    mlst.append(row[str(i)])
                self.metadata[row["Strain"]]["mlst_info"] = mlst
                metadata_count += 1

        # Go through the ROGA Summary file from the access DB to get strain/textual IDs, and 1' and 2' enzymes.
        # TODO: Make case sensitivity not matter here.
        try: # Assume we're using ROGA summary OLF. If it isn't there, assume ROGA summary OLC
            roga_summary = openpyxl.load_workbook("ROGA_summary_OLF.xlsx")
            # ws = roga_summary.get_sheet_by_name("ROGA_summary_OLF")
            ws = roga_summary.get_active_sheet()
            metadata_count = 0
            for row in ws.iter_rows(row_offset=1):
                if row[4].value in self.names:
                    self.metadata[row[4].value]["IsolateID"] = row[25].value
                    self.metadata[row[4].value]["TextualID"] = row[24].value
                    self.metadata[row[4].value]["1Enzyme"] = row[28].value
                    self.metadata[row[4].value]["2Enzyme"] = row[29].value
                    self.metadata[row[4].value]["Source"] = row[27].value
                    self.metadata[row[4].value]["ReceivedDate"] = row[1].value
                    self.metadata[row[4].value]["SequenceDate"] = row[2].value
                    self.metadata[row[4].value]["SequencedBy"] = row[3].value
                    metadata_count += 1

        except:  # Should be a file not found error - look it up.
            roga_summary = openpyxl.load_workbook("ROGA_summary_OLC.xlsx")
            # ws = roga_summary.get_sheet_by_name("ROGA_summary_OLF")
            ws = roga_summary.get_active_sheet()
            metadata_count = 0
            for row in ws.iter_rows(row_offset=1):
                if row[6].value in self.names:
                    self.metadata[row[6].value]["IsolateID"] = row[8].value
                    self.metadata[row[6].value]["TextualID"] = row[9].value
                    # self.metadata[row[6].value]["Source"] = row[27].value # Seemingly not necessary for VTEC, so that's good.
                    self.metadata[row[6].value]["ReceivedDate"] = row[10].value
                    self.metadata[row[6].value]["SequenceDate"] = row[11].value
                    self.metadata[row[6].value]["SequencedBy"] = row[12].value
                    metadata_count += 1
        # print(self.metadata)
        self.check_for_empty_data()

    # Look through resfinder data to see if AMR is present.
    def find_amr(self):
        """
        Looks through resfinder data to find any resistances present in any strains.
        Adds the resistances to the self.metadata dict, under the entry "AMR"
        """
        import csv
        amr_data = csv.DictReader(open(self.nasmnt + "WGSspades/reports/resfinder.csv"))
        for row in amr_data:
            name = row["Contig"].split("_")[0]
            if name in self.names:
                if float(row["PercentIdentity"]) > 98.0:
                    amr = list()
                    amr.append(row["Gene"])
                    amr.append(row["Resistance"])
                    amr.append(row["PercentIdentity"])
                    if "AMR" in self.metadata[name]:
                        self.metadata[name]["AMR"].append(amr)
                    else:
                        self.metadata[name]["AMR"] = list()
                        self.metadata[name]["AMR"].append(amr)

        # print(self.metadata)

    def add_amr_table(self):
        """
        Adds an AMR table to the document, if AMR was present, otherwise outputs the message the no AMR was detected.
        """
        import docx
        doc = docx.Document(self.outfile)
        # Check if any AMR is present within the samples. If it isn't, you don't have to add the AMR table.
        amr_present = False
        for name in self.names:
            if "AMR" in self.metadata[name]:
                amr_present = True
        para = doc.add_paragraph()
        para.add_run("Acquired Antimicrobial Resistance\n\n").bold = True
        if amr_present:
            para.add_run("Table 3:\n").bold = True
            amr_table = doc.add_table(1, 4)
            Automate.add_text_to_cell("Strain", 0, 0, amr_table, bold=True)
            Automate.add_text_to_cell("Gene", 0, 1, amr_table, bold=True)
            Automate.add_text_to_cell("Resistance", 0, 2, amr_table, bold=True)
            Automate.add_text_to_cell("Percent Identity", 0, 3, amr_table, bold=True)
            a = 0
            for name in self.names:
                if "AMR" in self.metadata[name]:
                    offset = a
                    for i in range(a, len(self.metadata[name]["AMR"]) + a):
                        amr_table.add_row()
                        Automate.add_text_to_cell(self.metadata[name]["TextualID"] + "\n(" + self.metadata[name]["IsolateID"] +
                                                  ")", i + 1, 0, amr_table)
                        a += 1
                        for j in range(3):
                            if j == 0:
                                Automate.add_text_to_cell(self.metadata[name]["AMR"][i - offset][j], i + 1, j + 1,
                                                          amr_table, italicize=True)
                            else:
                                Automate.add_text_to_cell(self.metadata[name]["AMR"][i - offset][j], i + 1, j + 1,
                                                          amr_table)
            para = doc.add_paragraph("AMR determination based on databases available at the Center for Genomic "
                                     "Epidimiology (")
            para.add_run("https://cge.cbs.dtu.dk/services/ResFinder/").underline = True
            para.add_run("). These tools currently focus on acquired genes, and therefore do not detect chromosomal "
                         "mutations.")
        else:
            para.add_run("No antimicrobial resistance genes were detected in these isolates.")

        doc.save(self.outfile)

    # Finds all seqIDs for sequences in the combined metadata file that match the MLST of any of the SEQIDs that we
    # are interested in for this ROGA.
    def find_matching_mlst(self):
        """
        Creates a dict (matching_mlst) that finds all seqIDs associated with any MLSTs present in the sample.
        Dict is in format {MLST1: [SeqID1, SeqID2, SeqID3], MLST2: [SeqID1, SeqID2]}
        """
        import csv
        metadata = csv.DictReader(open(self.nasmnt + "WGSspades/reports/combinedMetadata.csv"))
        for row in metadata:
            for name in self.names:
                if self.metadata[name]["MLST"] == row["MLSTsequencetype"]:
                    if self.metadata[name]["MLST"] not in self.matching_mlst:
                        self.matching_mlst[row["MLSTsequencetype"]] = [row["SampleName"]]
                    else:
                        if row["SampleName"] not in self.matching_mlst[row["MLSTsequencetype"]]:
                            self.matching_mlst[row["MLSTsequencetype"]].append(row["SampleName"])
        # print(self.matching_mlst)

    # Make a dict of MLST types.
    def mlst_to_dict(self):
        for strain in self.metadata:
            if self.metadata[strain]["MLST"] not in self.mlstdict:
                self.mlstdict[self.metadata[strain]["MLST"]] = [strain]
            else:
                self.mlstdict[self.metadata[strain]["MLST"]].append(strain)

    @staticmethod
    # This doesn't really work. Will need to do some more thinking on how to implement it.
    def merge_cells():
        import docx
        doc = docx.Document("modified.docx")
        tables = doc.tables
        for table in tables:
            for j in range(len(table.columns)):
                for i in range(len(table.rows) - 1):
                    if table.cell(i, j).text == table.cell(i + 1, j).text:
                        try:
                            # table.cell(i + 1, j).text = ""
                            table.cell(i, j).merge(table.cell(i + 1, j))
                        except:
                            pass
        doc.save("modified.docx")

    def add_redmine(self):
        """
        Creates a redmine request to do a SNVPhyl for any strains with non-unique MLSTs.
        Does this using the add_redmine_request.py script, which unfortunately has to be in python2
        and therefore can't be called in a better way than this.
        NOTE: To have this work, need to go to the add_redmine_request.py script and change the userid and password
        to actual credentials.
        """
        import subprocess
        for seqid in self.metadata:
            cmd = "/home/lowa/Virtual_Environments/web_stuff/bin/python add_redmine_request.py "
            reference = seqid
            compare = ""
            for item in self.matching_mlst[self.metadata[seqid]["MLST"]]:
                compare += item + ","
            compare = compare[:-1]

            cmd += reference + " " + compare
            process = subprocess.Popen(cmd.split(), stdout=subprocess.PIPE)
            output, error = process.communicate()
            # Output - will probably have to decode, because it comes out funny right now.
            print(output)

    def first_table(self):
        # TODO: Add in sequencing analyst, and OLC reference.
        # You could have saved yourself a lot of code by using datetime.date! Read docs better in the future.
        import docx
        investigator_dict = {"PM": "Paul Manninger", "Deschenesmy": "Mylene Deschenes", "CooperA": "Ashley Cooper",
                             "ManningerP": "Paul Manninger", "deschenesmy": "Mylene Deschenes", 'McMahonT': 'Tanis McMahon'}
        doc = docx.Document(self.outfile)
        table = doc.tables[0]
        # Create lists to keep track of what dates have already been received so duplicates aren't put in.
        received = list()
        sequenced = list()
        investigator = list()
        # Get the date received and date sequenced cells blanked out.
        table.cell(1, 2).text = ""
        table.cell(2, 2).text = ""
        Automate.add_text_to_cell("Adam Koziol", 1, 7, table)
        Automate.add_text_to_cell("Andrew Low/Catherine Carrillo", 2, 7, table)
        Automate.add_text_to_cell("Martine Gauthier/Burton Blais", 3, 7, table)
        # Iterate through all the strain names.
        for name in self.names:
            # Check that the received date hasn't already been put in.
            if self.metadata[name]["ReceivedDate"] not in received:
                # Add in the year.
                try:
                    Automate.add_text_to_cell(self.metadata[name]["ReceivedDate"].strftime("%Y-%m-%d"), 1, 2, table)
                except AttributeError:
                    Automate.add_text_to_cell(str(datetime.datetime.strptime(self.metadata[name]['ReceivedDate'], '%Y-%m-%d')), 1, 2, table)
                # Put date into list of dates already in there.
                received.append(self.metadata[name]["ReceivedDate"])
            # Exact same as for the received date, except for the sequenced date this time.
            if self.metadata[name]["SequenceDate"] not in sequenced:
                try:
                    Automate.add_text_to_cell(self.metadata[name]["SequenceDate"].strftime("%Y-%m-%d"), 2, 2, table)
                except AttributeError:
                    Automate.add_text_to_cell(str(datetime.datetime.strptime(self.metadata[name]['SequenceDate'], '%Y-%m-%d')), 2, 2, table)
                sequenced.append(self.metadata[name]["SequenceDate"])
            # Add in the sequencing analyst stuff.
            if self.metadata[name]["Investigator"] not in investigator:
                try:
                    Automate.add_text_to_cell(investigator_dict[self.metadata[name]["Investigator"]], 0, 7, table)
                except KeyError:
                    Automate.add_text_to_cell(self.metadata[name]["Investigator"], 0, 7, table)
                investigator.append(self.metadata[name]["Investigator"])
        # Finally, save the doc.
        doc.save(self.outfile)

    def find_recent_isolates(self, mlst, strain_name):
        # This actually doesn't ever get used any more. Can probably be removed.
        import openpyxl
        from datetime import datetime
        most_recent_strain = ""
        source = ""
        rdims = ""
        recent_received = ""
        shortest_diff = 1000000
        roga_summary = openpyxl.load_workbook("ROGA_summary_OLF.xlsx")
        ws = roga_summary.get_sheet_by_name("ROGA_summary_OLF")
        for row in ws.iter_rows(row_offset=1):
            if str(row[17].value) == mlst:
                received = row[6].value
                diff = self.day_diff(self.metadata[strain_name]["ReceivedDate"], received)
                if diff < shortest_diff and diff > 0 and str(row[7].value) not in self.metadata:
                    shortest_diff = diff
                    most_recent_strain = row[1].value
                    source = row[4].value
                    rdims = row[23].value
                    recent_received = row[6].value

        recent_received = datetime.date(recent_received)
        return most_recent_strain, source, rdims, recent_received

    def add_snvphyl_text(self, paragraph):
        import docx
        paragraph.add_run("\t\u2022 To determine the degree of relatedness to previous strains isolated at the CFIA "
                          "with the same sequence type, strains were further analyzed to identify single nucleotide"
                          " variants (SNVs). This approach compares the entire genome sequence to identify the number "
                          "of nucleotides that are different in at least one strain in the set of strains compared."
                          " SNVs were identified using the NML SNVPhyl pipeline (v1.0), using STRAIN_NAME as"
                          " the reference genome. (Biorequest# XXXX).\n")
        paragraph.add_run("\t\t\u2022 Maximum likelihood phylogeny based on XXX high quality core genome SNV positions"
                          " identified amongst YY genomes over ZZ.Z% of the reference genome is shown in ")
        paragraph.add_run(" Figure #\n").bold = True
        paragraph.add_run("\t\t\u2022 All of the new isolates have between X and Y SNVs compared to historical "
                         "isolates with the same genotype (")
        paragraph.add_run("Figure X, red boxes").bold = True
        paragraph.add_run(").\n\n")

    @staticmethod
    def day_diff(date1, date2):
        # from datetime import datetime
        # date1 = datetime.strptime(date1, "%Y-%M")
        # date2 = datetime.strptime(date2)
        try:
            return abs((date2 - date1).days)
        except TypeError:
            return 100000000

    def change_fonts(self):
        import docx
        from docx.shared import Pt
        doc = docx.Document(self.outfile)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    para = cell.paragraphs
                    for p in para:
                        for run in p.runs:
                            font = run.font
                            font.name = "Arial"
                            font.size = Pt(11)
        for para in doc.paragraphs:
            for run in para.runs:
                font = run.font
                font.name = "Arial"
                font.size = Pt(11)

        doc.save(self.outfile)

    @staticmethod
    def add_text_to_cell(text, x, y, table, italicize=False, bold=False):
        import docx
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        if italicize:
            table.cell(x, y).paragraphs[0].add_run(text).italic = True
        elif bold:
            table.cell(x, y).paragraphs[0].add_run(text).bold = True
        else:
            table.cell(x, y).paragraphs[0].add_run(text)
        table.cell(x, y).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # I have literally no idea how the next 5 lines of code work, but they seem to.
        # Taken from https://github.com/python-openxml/python-docx/issues/163
        tc = table.cell(x, y)._tc
        tcPr = tc.get_or_add_tcPr()
        tcVAlign = OxmlElement('w:vAlign')
        tcVAlign.set(qn('w:val'), "center")
        tcPr.append(tcVAlign)

    def check_for_empty_data(self):
        keys = ["Investigator", "Coverage", "TotalLength", "rST", "PipelineVersion", "MLST", "geneSeekr", "Matches"
                , "mlst_info", "IsolateID", "TextualID", "ReceivedDate", "SequenceDate",
                "SequencedBy"]
        for name in self.names:
            if name not in self.metadata:
                raise ROGAError("Isolate " + name + " not found in metadata!")
            for key in keys:
                try:
                    if self.metadata[name][key] is None:
                        raise ROGAError("Info " + key + " for isolate " + name + " was not found.")
                except KeyError:
                    raise ROGAError("Info " + key + " could not be found in the supplied data sheets.")


class ROGAError(ValueError):
    def __init__(self, message, *args):
        self.message = message  # without this you may get DeprecationWarning
        # allow users initialize misc. arguments as any other builtin Error
        super(ROGAError, self).__init__(message, *args)



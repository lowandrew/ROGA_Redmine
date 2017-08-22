# This is where I'll learn to use python-docx so that ROGAs can get automated.

from Automate import Automate


class VTEC_Automator(Automate):

    def get_serotype(self):
        import csv
        # Now we need to get the serotype since this is Ecoli.
        serotype_data = csv.DictReader(open(self.nasmnt + "WGSspades/reports/serotype.csv", encoding="ISO-8859-1"))
        for row in serotype_data:
            if row["Strain"] in self.names:
                serotype = dict()
                serotype["O"] = row["O-type"]
                serotype["H"] = row["H-type"]
                self.metadata[row["Strain"]]["serotype"] = serotype
        serotype_data = csv.DictReader(open(self.nasmnt + "External_WGSspades/reports/serotype.csv", encoding="ISO-8859-1"))
        for row in serotype_data:
            if row["Strain"] in self.names:
                serotype = dict()
                serotype["O"] = row["O-type"]
                serotype["H"] = row["H-type"]
                self.metadata[row["Strain"]]["serotype"] = serotype

    # This prints the GeneSeekr Analysis table and the Sequence Data Quality Table to the doc.
    def print_tables_to_docx(self):
        import docx
        doc = docx.Document("vtec_template.docx")
        j = 2
        # GeneSeekr Analysis Table.
        for name in self.names:
            # Add a new row to the table for each SEQID to be added.
            doc.tables[2].add_row()
            # Input sample names.
            Automate.add_text_to_cell(self.metadata[name]["TextualID"] + "\n("
                                "" + self.metadata[name]["IsolateID"] + ")", j, 0, doc.tables[2])
            # Input uidA presence/absence
            if "uidA" in self.metadata[name]["geneSeekr"]:
                Automate.add_text_to_cell(u"\u2022", j, 1, doc.tables[2])
            else:
                Automate.add_text_to_cell("ND", j, 1, doc.tables[2])
            # Input O and H types.
            Automate.add_text_to_cell(self.metadata[name]["serotype"]["O"].replace("(100.0)", ""), j, 2, doc.tables[2])
            Automate.add_text_to_cell(self.metadata[name]["serotype"]["H"].replace("(100.0)", ""), j, 3, doc.tables[2])
            # Input the Ecoli specific virulence genes.
            # This way of doing things is really inflexible and generally not great, but it'll work for now.
            if "hlyA" in self.metadata[name]["geneSeekr"]:
                Automate.add_text_to_cell(u"\u2022", j, 4, doc.tables[2])
            else:
                Automate.add_text_to_cell("ND", j, 4, doc.tables[2])
            # The following stretch of code is amongst the ugliest stuff you've ever written.
            # Output eae genes
            present = False
            for item in self.metadata[name]["geneSeekr"]:
                if "eae" in item:
                    Automate.add_text_to_cell(item + ";", j, 5, doc.tables[2], italicize=True)
                    present = True
            if not present:
                Automate.add_text_to_cell("ND", j, 5, doc.tables[2])
            # Output vt1 genes.
            present = False
            for item in self.metadata[name]["geneSeekr"]:
                if "vt1" in item or "VT1" in item:
                    Automate.add_text_to_cell(item + ";", j, 6, doc.tables[2], italicize=True)
                    present = True
            if not present:
                Automate.add_text_to_cell("ND", j, 5, doc.tables[2])
            # Output vt2 genes.
            present = False
            for item in self.metadata[name]["geneSeekr"]:
                if "vt2" in item or "VT2" in item:
                    Automate.add_text_to_cell(item + ";", j, 7, doc.tables[2], italicize=True)
                    present = True
            if not present:
                Automate.add_text_to_cell("ND", j, 7, doc.tables[2])
            # Input MLST types.
            # Enter the MLST type.
            Automate.add_text_to_cell(self.metadata[name]["MLST"], j, 9, doc.tables[2])
            # Enter the type for each of the seven genes. Again, this code is inflexible and not great, but works.
            for i in range(10, 17):
                Automate.add_text_to_cell(self.metadata[name]["mlst_info"][i - 10], j, i, doc.tables[2])
            doc.tables[2].cell(j, 8).merge(doc.tables[2].cell(j, 9))
            Automate.add_text_to_cell("", j, 9, doc.tables[2])
            j += 1

        # Sequence Data quality table.
        j = 1
        for name in self.names:
            # Add a new row to the table for each SEQID to be added.
            doc.tables[3].add_row()
            # Input SEQIDS.
            Automate.add_text_to_cell(self.metadata[name]["TextualID"] + "\n"
                                        "(" + self.metadata[name]["IsolateID"] + ")", j, 0, doc.tables[3])
            Automate.add_text_to_cell(name, j, 1, doc.tables[3])
            doc.tables[3].cell(j, 1).text = name
            # Input total length.
            length = int(self.metadata[name]["TotalLength"])
            Automate.add_text_to_cell(str(length), j, 2, doc.tables[3])
            # Input coverage.
            cov = float(self.metadata[name]["Coverage"])
            covstr = "%.1f" % cov
            Automate.add_text_to_cell(covstr, j, 3, doc.tables[3])
            # Input % Identity GDCS, rST, and Pass/Fail.
            percent_id = float(self.metadata[name]["Matches"])/53.0 * 100.0
            Automate.add_text_to_cell(str(percent_id), j, 4, doc.tables[3])
            Automate.add_text_to_cell(self.metadata[name]["rST"], j, 5, doc.tables[3])
            if percent_id == 100.0:
                Automate.add_text_to_cell("Pass", j, 6, doc.tables[3])
            else:
                Automate.add_text_to_cell("Fail", j, 6, doc.tables[3])
            # Finally, input pipeline version.
            Automate.add_text_to_cell(self.metadata[name]["PipelineVersion"], j, 7, doc.tables[3])
            j += 1
        doc.save(self.outfile)

    def find_verotoxin_genes(self):
        verotoxin_genes = list()
        for strain in self.metadata:
            for item in self.metadata[strain]["geneSeekr"]:
                if "VT" in item or "vt" in item:
                    if item not in verotoxin_genes:
                        verotoxin_genes.append(item)
        return verotoxin_genes

    # This takes care of printing the SNVPhyl summary to docx.
    def print_summary_to_docx(self):
        import docx
        import re
        doc = docx.Document(self.outfile)
        tables = doc.tables
        # This is going to be long and messy, but I'm not sure of a better way to do it.
        tables[1].cell(0, 0).text = ""
        para = tables[1].cell(0, 0).add_paragraph("")
        para.add_run("Identification Summary:\n\n").bold = True
        if len(self.names) == 1:
            para.add_run("Strain " + self.metadata[self.names[0]]["TextualID"])
            para.add_run(" was submitted to for whole-genome sequencing and confirmed to be ")
        else:
            para.add_run(str(len(self.names)) + " strains (see Table 1)")
            para.add_run(" were submitted for whole-genome sequencing and confirmed to be ")
        para.add_run("VTEC")
        para.add_run(" based on the detection of probe sequences (e-probes) indicating the presence of ")
        para.add_run("verotoxin")
        vt_genes = self.find_verotoxin_genes()
        if len(vt_genes) == 1:
            para.add_run(" gene " + vt_genes[0] + ".\n\n")
        else:
            para.add_run(" genes ")
            for i in range(len(vt_genes)):
                if i == len(vt_genes) - 2:
                    para.add_run(vt_genes[i] + " and ")
                elif i == len(vt_genes) - 1:
                    para.add_run(vt_genes[i] + ", ")
                else:
                    para.add_run(vt_genes[i] + ", ")
        para.add_run("and intimin (")
        para.add_run("eae").italic = True
        para.add_run(") genes.\n\n")
        if len(self.metadata) == 1:
            para.add_run("Further analyses conducted using databases from the Center For Genomic Epidemiology"
                     " (https://cge.cbs.dtu.dk/services/SerotypeFinder/) predicted the serotype of this isolate to be ")
            for strain in self.metadata:
                para.add_run(re.sub(r'\([^)]*\)', '', self.metadata[strain]["serotype"]["O"])
                             + ":" + re.sub(r'\([^)]*\)', '', self.metadata[strain]["serotype"]["H"]) + ".\n")
        else:
            para.add_run("Further analyses conducted using databases from the Center For Genomic Epidemiology"
                         " (https://cge.cbs.dtu.dk/services/SerotypeFinder/) predicted the serotypes of these isolates:\n")
            for strain in self.metadata:
                para.add_run("\t\u2022Isolate " + self.metadata[strain]["TextualID"] + " was serotype " + re.sub(r'\([^)]*\)', '', self.metadata[strain]["serotype"]["O"])
                             + ":" + re.sub(r'\([^)]*\)', '', self.metadata[strain]["serotype"]["H"]) + "\n")
        para.add_run("\nMultilocus sequence typing (MLST):\n\n").underline = True
        if len(self.mlstdict) > 1:
            para.add_run("\u2022 Isolates were sequence types ")
            i = 1
            for st in self.mlstdict:
                para.add_run("ST-")
                para.add_run(st)
                if i == len(self.mlstdict) - 1:
                    para.add_run(" and ")
                elif i == len(self.mlstdict):
                    para.add_run(".\n")
                else:
                    para.add_run(", ")
                i += 1
        else:
            para.add_run("\u2022 All isolates were sequence type ")
            for st in self.mlstdict:
                para.add_run("ST-")
                para.add_run(st + ".\n")

        for st in self.matching_mlst:
            if len(self.matching_mlst[st]) > 5:
                para.add_run("\tIsolates with ST-" + st + " are commonly recovered from the CFIA's food testing"
                                                          " program (2009-present).\n")
            elif len(self.matching_mlst[st]) == 1:
                para.add_run("\tIsolates with ST-" + st + " have not previously been recovered by the CFIA's"
                                                          " food testing program (2009-present).\n")
            else:
                para.add_run("\tIsolates with ST-" + st + " have previously been recovered from the CFIA's food "
                                                                  "testing program (2009-present).\n")

        para.add_run("\n")
        para.add_run("Quality Control Analysis:\n").bold = True
        if "11" not in self.matching_mlst:
            if len(self.metadata) == 1:
                para.add_run("\t This isolate does not match CFIA ")
            else:
                para.add_run("\t These isolates do not match the ")
            para.add_run("OLC-795 nalidixic acid-resistant ")
            para.add_run("E. coli").italic = True
            para.add_run(" control strain used at the CFIA (ST-11, rST-2119).")
        else:
            matches_control = False
            for strain in self.metadata:
                if self.metadata[strain]["MLST"] == 11 and self.metadata[strain]["rST"] == 2119:
                    matches_control = True
                    para.add_run("Isolate " + self.metadata[strain]["TextualID"] + " matches")
                    para.add_run(" the OLC-795 nalidixic acid-resistant ")
                    para.add_run("E. coli").italic = True
                    para.add_run(" control strain used at the CFIA (ST-11, rST-2119).")
            if not matches_control:
                if len(self.metadata) == 1:
                    para.add_run("\t This isolate does not match CFIA ")
                else:
                    para.add_run("\t These isolates do not match the ")
            para.add_run("OLC-795 nalidixic acid-resistant ")
            para.add_run("E. coli").italic = True
            para.add_run(" control strain used at the CFIA (ST-11, rST-2119).")


        doc.save(self.outfile)

    def print_references(self):
        import docx
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = docx.Document(self.outfile)
        para = doc.add_paragraph("")
        para.add_run("\n\nReferences:\n").bold = True
        para.add_run("\t1.  Hayashi T, Makino K, Ohnishi M, Kurokawa K, Ishii K, Yokoyama K, Han CG, Ohtsubo E, "
                     "Nakayama K, Murata T ")
        para.add_run("et al:").italic = True
        para.add_run(" Complete genome sequence of enterohemorrhagic Eshcherichia coli O157:H7 and genomic comparison"
                     " with laboratory strain K-12.").bold = True
        para.add_run(" DNA Res. 2001; 8(1):11-22\n")
        para.add_run("\t2.  Jaureguy F, Landraud L, Passet V, Diancourt L, Frapy E, Guigon G, et al. ")
        para.add_run("Phylogenetic and genomic diversity of human bateremic Escherichia coli strains.").bold = True
        para.add_run(" BMC Genomics. 2008; 9:560.\n")
        para.add_run("\t3.  Jolley KA, Bliss CM, Bennett JS, Bratcher HB, et al. (2012) ")
        para.add_run("Ribosomal multilocus sequence typing: universal characterization of bacteria from domain to "
                     "strain.").bold = True
        para.add_run(" Microbiology. 158:1005-15")
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run("--End Of Report--").bold = True
        doc.save(self.outfile)

    def __init__(self, args):
        self.seq_id_list = args.seq_id_list
        self.outfile = args.outfile_name
        self.nasmnt = args.nasmnt
        self.names = dict()
        self.metadata = dict()
        self.matching_mlst = dict()
        self.parse_metadata()
        self.get_serotype()
        self.find_matching_mlst()
        self.mlstdict = dict()
        self.mlst_to_dict()
        self.print_tables_to_docx()
        self.print_summary_to_docx()
        self.first_table()
        self.find_amr()
        self.add_amr_table()
        self.print_references()
        self.change_fonts()
        # self.add_redmine()


if __name__ == "__main__":

    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("seq_id_list")
    parser.add_argument("outfile_name")
    parser.add_argument("-n", "--nasmnt", type=str, default="/mnt/nas/", help="Where your NAS is mounted. Default is "
                                                                              "/mnt/nas/, trailing slash must be"
                                                                              "included.")
    arguments = parser.parse_args()
    x = VTEC_Automator(arguments)

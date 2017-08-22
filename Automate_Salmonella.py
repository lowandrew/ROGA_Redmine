# This is where I'll learn to use python-docx so that ROGAs can get automated.

from Automate import Automate


class Salmonella_Automator(Automate):

    def get_sistr_data(self):
        import csv
        with open(self.nasmnt + "WGSspades/reports/sistr.tsv") as sis:
            sistr_data = csv.DictReader(sis, delimiter="\t")
            for row in sistr_data:
                if row["genome"].replace(".fasta", "") in self.names:
                    self.metadata[row["genome"].replace(".fasta","")]["serovar"] = row["serovar"]
                    self.metadata[row["genome"].replace(".fasta","")]["h1"] = row["h1"]
                    self.metadata[row["genome"].replace(".fasta","")]["h2"] = row["h2"]
                    self.metadata[row["genome"].replace(".fasta","")]["serogroup"] = row["serogroup"]
        with open(self.nasmnt + "External_WGSspades/reports/sistr.tsv") as sis:
            sistr_data = csv.DictReader(sis, delimiter="\t")
            for row in sistr_data:
                if row["genome"].replace(".fasta", "") in self.names:
                    self.metadata[row["genome"].replace(".fasta","")]["serovar"] = row["serovar"]
                    self.metadata[row["genome"].replace(".fasta","")]["h1"] = row["h1"]
                    self.metadata[row["genome"].replace(".fasta","")]["h2"] = row["h2"]
                    self.metadata[row["genome"].replace(".fasta","")]["serogroup"] = row["serogroup"]



    # This prints the GeneSeekr Analysis table and the Sequence Data Quality Table to the doc.
    def print_tables_to_docx(self):
        import docx
        doc = docx.Document("salmonella_template.docx")
        j = 2
        # GeneSeekr Analysis Table.
        for name in self.names:
            # Add a new row to the table for each SEQID to be added.
            doc.tables[2].add_row()
            # Input sample names.
            Automate.add_text_to_cell(self.metadata[name]["TextualID"] + ""
                        "\n(" + self.metadata[name]["IsolateID"] + ")", j, 0, doc.tables[2])
            Automate.add_text_to_cell(self.metadata[name]["serovar"], j, 1, doc.tables[2])
            Automate.add_text_to_cell(self.metadata[name]["Source"], j, 2, doc.tables[2])
            Automate.add_text_to_cell(self.metadata[name]["1Enzyme"], j, 3, doc.tables[2])
            Automate.add_text_to_cell(self.metadata[name]["2Enzyme"], j, 4, doc.tables[2])
            # Input the salmonella specific marker genes.
            # This way of doing things is really inflexible and generally not great, but it'll work for now.
            present = False
            for item in self.metadata[name]["geneSeekr"]:
                if "INVA" in item.upper():
                    present = True
            if not present:
                Automate.add_text_to_cell("ND", j, 5, doc.tables[2])
            else:
                Automate.add_text_to_cell(u"\u2022", j, 5, doc.tables[2])
            present = False
            for item in self.metadata[name]["geneSeekr"]:
                if "STN" in item.upper():
                    present = True
            if not present:
                Automate.add_text_to_cell("ND", j, 6, doc.tables[2])
            else:
                Automate.add_text_to_cell(u"\u2022", j, 6, doc.tables[2])
            # Input MLST types.
            # Enter the MLST type.
            Automate.add_text_to_cell(self.metadata[name]["MLST"], j, 7, doc.tables[2])
            # Enter the type for each of the seven genes. Again, this code is inflexible and not great, but works.
            for i in range(8, 15):
                Automate.add_text_to_cell(self.metadata[name]["mlst_info"][i - 8], j, i, doc.tables[2])
            j += 1

        # Sequence Data quality table.
        j = 1
        for name in self.names:
            # Add a new row to the table for each SEQID to be added.
            doc.tables[3].add_row()
            # Input SEQIDS.
            # TODO HERE: Figure out where the sample names are stored and put them into the first column.
            Automate.add_text_to_cell(self.metadata[name]["TextualID"] + ""
                                      "\n(" + self.metadata[name]["IsolateID"] + ")", j, 0, doc.tables[3])
            Automate.add_text_to_cell(name, j, 1, doc.tables[3])
            # Input total length.
            length = int(self.metadata[name]["TotalLength"])
            Automate.add_text_to_cell(str(length), j, 2, doc.tables[3])
            # Input coverage.
            cov = float(self.metadata[name]["Coverage"])
            covstr = "%.1f" % cov
            Automate.add_text_to_cell(covstr, j, 3, doc.tables[3])
            # Input % Identity GDCS, rST, and Pass/Fail.
            percent_id = float(self.metadata[name]["Matches"])/53.0 * 100.0
            Automate.add_text_to_cell(str(percent_id), j, 4, doc.tables[3])
            # Input % Identity GDCS, rST, and Pass/Fail.
            Automate.add_text_to_cell(self.metadata[name]["rST"], j, 5, doc.tables[3])
            if percent_id == 100.0:
                Automate.add_text_to_cell("Pass", j, 6, doc.tables[3])
            else:
                Automate.add_text_to_cell("Fail", j, 6, doc.tables[3])
            # Finally, input pipeline version.
            Automate.add_text_to_cell(self.metadata[name]["PipelineVersion"], j, 7, doc.tables[3])
            j += 1
        doc.save(self.outfile)

    # This takes care of printing the SNVPhyl summary to docx.
    def print_summary_to_docx(self):
        import docx
        doc = docx.Document(self.outfile)
        tables = doc.tables
        # This is going to be long and messy, but I'm not sure of a better way to do it.
        tables[1].cell(0, 0).text = ""
        para = tables[1].cell(0, 0).add_paragraph("")
        para.add_run("Identification Summary:\n\n").bold = True
        if len(self.names) == 1:
            para.add_run("Strain " + self.metadata[self.names[0]]["TextualID"] + " was submitted for whole-genome sequencing and confirmed to be ")
        else:
            para.add_run(str(len(self.names)) + " strains (see Table 1)")
            para.add_run(" were submitted for whole-genome sequencing and confirmed to be ")
        para.add_run("Salmonella").italic = True
        para.add_run(" based on the detection of probe sequences (e-probes) indicating the presence of ")
        para.add_run("invA").italic = True
        para.add_run(" and ")
        para.add_run("stn").italic = True
        para.add_run(" genes.\n\n")
        para.add_run("Multilocus sequence typing (MLST):\n\n").underline = True
        if len(self.mlstdict) > 1:
            para.add_run("\u2022 Isolates were sequence types ")
            i = 1
            for st in self.mlstdict:
                para.add_run("ST-")
                para.add_run(st)
                if i == len(self.mlstdict) - 1:
                    para.add_run(" and ")
                elif i == len(self.mlstdict):
                    para.add_run(".\n")
                else:
                    para.add_run(", ")
                i += 1
        else:
            para.add_run("\u2022 All isolates were sequence type ")
            for st in self.mlstdict:
                para.add_run("ST-")
                para.add_run(st + ".\n")

        for st in self.matching_mlst:
            if len(self.matching_mlst[st]) > 5:
                para.add_run("\tIsolates with ST-" + st + " are commonly recovered from the CFIA's food testing"
                                                          " program (2009-present).\n")
            elif len(self.matching_mlst[st]) == 1:
                para.add_run("\tIsolates with ST-" + st + " have not previously been recovered by the CFIA's"
                                                          " food testing program (2009-present).\n")
            else:
                para.add_run("\tIsolates with ST-" + st + " have previously been recovered from the CFIA's food "
                                                                  "testing program (2009-present).\n")

        para.add_run("\n")
        para.add_run("Salmonella In Silico Typing Resource").italic = True
        para.add_run("(SISTR; https://lfz.corefacility.ca/sistr-app/ (PHAC)) analysis[1] (Identification matching "
                     "phenotype is highlighted in bold):\n\n")
        for name in self.names:
            para.add_run("\t\u2022  Strain " + self.metadata[name]["TextualID"] + " was predicted to be serovar ")
            para.add_run(self.metadata[name]["serovar"] + "\n").bold = True

        para.add_run("\nQuality Control Analysis:\n\n").bold = True
        if "2666" not in self.matching_mlst:
            if len(self.metadata) == 1:
                para.add_run("\t\u2022 This isolate des not match the ")
            else:
                para.add_run("\t\u2022 These isolates do not match the ")
            para.add_run("Salmonella").italic = True
            para.add_run(" ser. Mishmarhaemek (ST-2666) control strain used in CFIA food testing laboratories.")
        else:
            for strain in self.metadata:
                if self.metadata[strain]["MLST"] == "2666":
                    para.add_run("\t\u2022 Isolate " + self.metadata[strain]["TextualID"] + " matches the ")
                    para.add_run("Salmonella").italic = True
                    para.add_run(" ser. Mishmarhaemek (ST-2666) control strain used in CFIA food testing laboratories.")

        doc.save(self.outfile)

    def print_references(self):
        import docx
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = docx.Document(self.outfile)
        para = doc.add_paragraph("")
        para.add_run("\n\nReferences:\n").bold = True
        para.add_run("\t1.  Yoshida C, Kruczkiewicz P, Laing CR, Lingohr EJ, Gannon VPJ, Nash JHE, Taboada EN. The")
        para.add_run(" Salmonella In Silico").italic = True
        para.add_run(" Typing Resource (SISTR): an open web-accessible tool for rapidly tping and subtpying draft "
                     "Salmonella genome assemblies. PLoS ONE 11(1):e0147101. doi: 10.1371/journal.pone.0147101\n")
        para.add_run("\t2.  Achtman M, Wain J, Weill FX, Nair S, Zhou Z, et al. (2012) Multilocus sequence typing as a"
                     "replacement for serotyping in ")
        para.add_run("Salmonella enterica").italic = True
        para.add_run(". PLoS Pathog 8: e1002776.\n")
        para.add_run("\t3.  Jolley KA, Bliss CM, Bennett JS, Bratcher HB, et al. (2012) Ribosomal multilocus sequence"
                     " typing: universal characterization of bacteria from domain to strain. Microbiology. 158:1005-15")
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run("--End Of Report--").bold = True
        doc.save(self.outfile)

    def __init__(self, args):
        self.seq_id_list = args.seq_id_list
        self.outfile = args.outfile_name
        self.nasmnt = args.nasmnt
        self.names = dict()
        self.metadata = dict()
        self.matching_mlst = dict()
        self.mlstdict = dict()
        self.parse_metadata()
        self.get_sistr_data()
        self.find_matching_mlst()
        self.mlst_to_dict()
        self.print_tables_to_docx()
        self.print_summary_to_docx()
        self.find_amr()
        self.add_amr_table()
        self.first_table()
        # Automate.merge_cells()
        # self.add_redmine()
        self.print_references()
        self.change_fonts()

if __name__ == "__main__":

    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("seq_id_list")
    parser.add_argument("outfile_name")
    parser.add_argument("-n", "--nasmnt", type=str, default="/mnt/nas/", help="Where your NAS is mounted. Default is "
                                                                              "/mnt/nas/, trailing slash must be"
                                                                              "included.")
    arguments = parser.parse_args()
    x = Salmonella_Automator(arguments)
